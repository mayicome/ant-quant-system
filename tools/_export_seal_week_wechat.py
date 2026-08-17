# -*- coding: utf-8 -*-
"""导出封单评级周复盘：公众号用 PNG + Word。"""
from __future__ import annotations

from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.gridspec import GridSpec
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm, Inches

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "history_data"
PNG_PATH = OUT_DIR / "封单评级验证_上周四至这周四_公众号.png"
DOCX_PATH = OUT_DIR / "封单评级验证_上周四至这周四_公众号.docx"

TITLE = "封单评级周复盘：上周四至这周四"
SUB = "封板日 7/30–8/06 · 次日验证 · 有效样本 540 只 · 8/07 封板待下周一验证未纳入"

POOL = [
    {"tier": "强", "n": 50, "open": 7.68, "close": 7.63, "low": 3.89, "gap": 0.0, "up": 88.0, "hold": -0.05},
    {"tier": "中", "n": 202, "open": 2.19, "close": 3.26, "low": -1.40, "gap": 14.4, "up": 67.8, "hold": 1.07},
    {"tier": "弱", "n": 288, "open": 0.54, "close": 2.83, "low": -2.25, "gap": 28.8, "up": 53.5, "hold": 2.29},
]

DAILY = [
    ("07/30", "07/31", "参考", "0/0/51", "— / — / -0.44%", "无法分层（全弱）"),
    ("07/31", "08/03", "有效", "7/35/54", "+6.01 / +2.32 / +1.58%", "+4.43pp"),
    ("08/03", "08/04", "有效", "11/36/28", "+9.83 / +3.73 / +1.49%", "+8.34pp"),
    ("08/04", "08/05", "有效", "10/45/82", "+7.24 / +1.72 / +0.28%", "+6.96pp"),
    ("08/05", "08/06", "有效", "12/46/44", "+6.77 / +1.72 / +0.87%", "+5.90pp"),
    ("08/06", "08/07", "有效", "10/40/29", "+8.02 / +1.75 / -0.31%", "+8.33pp"),
]

HEADLINE = (
    "可分层的 5 个交易日全部判「分层有效」。"
    "强档开盘均 +7.68%，弱档仅 +0.54%，深低开占比 0% vs 28.8%。"
    "优势主要在开盘质量，不是多拿到尾盘。"
)

POINTS = [
    "窗口：封板日上周四（7/30）至这周四（8/06）；周五封板需下周一验证，未纳入。",
    "样本：有效 540 只；强 50 / 中 202 / 弱 288。强档日均约 10 只，足够对照。",
    "开盘分层清晰：强→中→弱 开盘均 +7.68% / +2.19% / +0.54%，高开占比 88% / 68% / 54%。",
    "风险分层清晰：深低开（开盘<-1%）占比 0% / 14.4% / 28.8%，强档全程零深低开。",
    "可分层日 5/5 有效；仅 7/30 全为弱档、无法分层，不影响其余结论。",
    "持有到收盘不是额外 Alpha：强档「持有相对开盘卖」周均约 0；弱/中为正，多为低开后修复。",
]


def _setup_font() -> None:
    matplotlib.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "Arial Unicode MS", "sans-serif"]
    matplotlib.rcParams["axes.unicode_minus"] = False


def export_png(path: Path) -> Path:
    _setup_font()
    # 公众号长图：宽约 1080px @2x ≈ 5.4 inch @200dpi
    fig = plt.figure(figsize=(5.4, 11.2), facecolor="#FFFFFF", dpi=200)
    gs = GridSpec(5, 1, height_ratios=[0.9, 0.7, 1.6, 1.5, 2.2], hspace=0.35, left=0.08, right=0.95, top=0.97, bottom=0.03)

    ax0 = fig.add_subplot(gs[0])
    ax0.axis("off")
    ax0.text(0.0, 0.72, TITLE, fontsize=15, fontweight="bold", color="#111111", transform=ax0.transAxes, va="top")
    ax0.text(0.0, 0.28, SUB, fontsize=8.5, color="#666666", transform=ax0.transAxes, va="top")

    ax1 = fig.add_subplot(gs[1])
    ax1.axis("off")
    ax1.set_xlim(0, 1)
    ax1.set_ylim(0, 1)
    ax1.add_patch(plt.Rectangle((0, 0.05), 1, 0.9, facecolor="#E8F5E9", edgecolor="#A5D6A7", linewidth=0.8, transform=ax1.transAxes))
    ax1.text(0.03, 0.55, HEADLINE, fontsize=9, color="#1B5E20", transform=ax1.transAxes, va="center", wrap=True)

    # 分档表
    ax2 = fig.add_subplot(gs[2])
    ax2.axis("off")
    ax2.text(0.0, 1.02, "分档周聚合", fontsize=11, fontweight="bold", transform=ax2.transAxes, va="bottom")
    cols = ["评级", "只数", "开盘均%", "收盘均%", "最低均%", "低开占比%", "高开占比%"]
    cells = []
    for r in POOL:
        cells.append(
            [
                r["tier"],
                str(r["n"]),
                f'{r["open"]:+.2f}%',
                f'{r["close"]:+.2f}%',
                f'{r["low"]:+.2f}%',
                f'{r["gap"]:.1f}%',
                f'{r["up"]:.1f}%',
            ]
        )
    table = ax2.table(cellText=cells, colLabels=cols, loc="upper center", cellLoc="center")
    table.auto_set_font_size(False)
    table.set_fontsize(8)
    table.scale(1.0, 1.55)
    colors = {"强": "#E8F5E9", "中": "#FFF8E1", "弱": "#FFEBEE"}
    for (r, c), cell in table.get_celld().items():
        cell.set_edgecolor("#DDDDDD")
        if r == 0:
            cell.set_facecolor("#F0F4F8")
            cell.set_text_props(weight="bold")
        elif r > 0:
            tier = cells[r - 1][0]
            if c == 0:
                cell.set_facecolor(colors.get(tier, "#FFFFFF"))

    # 柱状图
    ax3 = fig.add_subplot(gs[3])
    tiers = [r["tier"] for r in POOL]
    opens = [r["open"] for r in POOL]
    gaps = [r["gap"] for r in POOL]
    x = range(len(tiers))
    w = 0.35
    bars1 = ax3.bar([i - w / 2 for i in x], opens, width=w, color="#2E7D32", label="开盘均%")
    bars2 = ax3.bar([i + w / 2 for i in x], gaps, width=w, color="#C62828", label="低开占比%")
    ax3.set_xticks(list(x))
    ax3.set_xticklabels(tiers)
    ax3.axhline(0, color="#999999", linewidth=0.6)
    ax3.legend(fontsize=7, loc="upper right", frameon=False)
    ax3.set_title("开盘均% vs 低开占比%", fontsize=10, pad=6)
    ax3.tick_params(labelsize=8)
    for spine in ("top", "right"):
        ax3.spines[spine].set_visible(False)

    # 每日表
    ax4 = fig.add_subplot(gs[4])
    ax4.axis("off")
    ax4.text(0.0, 1.02, "每日摘要", fontsize=11, fontweight="bold", transform=ax4.transAxes, va="bottom")
    dcols = ["封板日", "验证日", "结论", "强/中/弱", "开盘 强/中/弱", "强−弱差"]
    dcells = [list(row) for row in DAILY]
    t2 = ax4.table(cellText=dcells, colLabels=dcols, loc="upper center", cellLoc="center")
    t2.auto_set_font_size(False)
    t2.set_fontsize(7.2)
    t2.scale(1.0, 1.35)
    for (r, c), cell in t2.get_celld().items():
        cell.set_edgecolor("#DDDDDD")
        if r == 0:
            cell.set_facecolor("#F0F4F8")
            cell.set_text_props(weight="bold")
        elif r > 0 and c == 2:
            v = dcells[r - 1][2]
            if v == "有效":
                cell.set_facecolor("#E8F5E9")
            elif v == "参考":
                cell.set_facecolor("#F5F5F5")

    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white", pad_inches=0.12)
    plt.close(fig)
    return path


def _set_run_font(run, name="微软雅黑", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def export_docx(path: Path, png_path: Path) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    _set_run_font(r, size=18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SUB)
    _set_run_font(r, size=9, color=RGBColor(0x66, 0x66, 0x66))

    p = doc.add_paragraph()
    r = p.add_run("一句话结论")
    _set_run_font(r, size=13, bold=True)
    p = doc.add_paragraph()
    r = p.add_run(HEADLINE)
    _set_run_font(r, size=11)

    p = doc.add_paragraph()
    r = p.add_run("核心数据")
    _set_run_font(r, size=13, bold=True)

    table = doc.add_table(rows=1 + len(POOL), cols=7)
    table.style = "Table Grid"
    headers = ["评级", "只数", "开盘均%", "收盘均%", "最低均%", "低开占比%", "高开占比%"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        rr = cell.paragraphs[0].add_run(h)
        _set_run_font(rr, size=9, bold=True)
    for i, row in enumerate(POOL):
        vals = [
            row["tier"],
            str(row["n"]),
            f'{row["open"]:+.2f}%',
            f'{row["close"]:+.2f}%',
            f'{row["low"]:+.2f}%',
            f'{row["gap"]:.1f}%',
            f'{row["up"]:.1f}%',
        ]
        for j, v in enumerate(vals):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            rr = cell.paragraphs[0].add_run(v)
            _set_run_font(rr, size=9)

    p = doc.add_paragraph()
    r = p.add_run("每日摘要")
    _set_run_font(r, size=13, bold=True)
    dtable = doc.add_table(rows=1 + len(DAILY), cols=6)
    dtable.style = "Table Grid"
    dh = ["封板日", "验证日", "结论", "强/中/弱", "开盘强/中/弱", "强−弱差"]
    for i, h in enumerate(dh):
        cell = dtable.rows[0].cells[i]
        cell.text = ""
        rr = cell.paragraphs[0].add_run(h)
        _set_run_font(rr, size=9, bold=True)
    for i, row in enumerate(DAILY):
        for j, v in enumerate(row):
            cell = dtable.rows[i + 1].cells[j]
            cell.text = ""
            rr = cell.paragraphs[0].add_run(v)
            _set_run_font(rr, size=8)

    p = doc.add_paragraph()
    r = p.add_run("怎么读")
    _set_run_font(r, size=13, bold=True)
    for i, line in enumerate(POINTS, 1):
        p = doc.add_paragraph()
        rr = p.add_run(f"{i}. {line}")
        _set_run_font(rr, size=10.5)

    if png_path.is_file():
        p = doc.add_paragraph()
        r = p.add_run("配图（可直接用于公众号）")
        _set_run_font(r, size=13, bold=True)
        doc.add_picture(str(png_path), width=Inches(5.8))

    p = doc.add_paragraph()
    r = p.add_run("说明：数据来自每日「封单评级验证日报」；评级按 v2 规则。公众号可直接用配图，或复制正文表格。")
    _set_run_font(r, size=8, color=RGBColor(0x88, 0x88, 0x88))

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def main() -> None:
    png = export_png(PNG_PATH)
    docx = export_docx(DOCX_PATH, png)
    print(f"PNG: {png}")
    print(f"DOCX: {docx}")


if __name__ == "__main__":
    main()
